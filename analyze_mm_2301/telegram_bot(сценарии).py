import asyncio
import logging
from aiogram import Bot, Dispatcher, types
from aiogram.filters import Command
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
import os
from datetime import datetime
from scenario_checker import check_meeting_scenario, safe_telegram_message

# Импортируем модуль для проверки критериев успеха
from success_criteria import check_success_criteria

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)

# Конфигурация бота
BOT_TOKEN = "8079592721:AAGLaX7LwUPX0X5fr1SK-9IQnSvfP3Z96ws"
bot = Bot(token=BOT_TOKEN)
storage = MemoryStorage()
dp = Dispatcher(storage=storage)


# Состояния FSM
class UserStates(StatesGroup):
    waiting_for_scenario_check = State()
    waiting_for_success_criteria = State()  # Новое состояние для проверки критериев успеха


def escape_telegram_chars(text: str) -> str:
    """Экранирование специальных символов для Telegram"""
    escape_chars = ['_', '*', '[', ']', '(', ')', '~', '`', '>', '#', '+', '-', '=', '|', '{', '}', '.', '!']
    for char in escape_chars:
        text = text.replace(char, f'\\{char}')
    return text


def filter_facilitator_speech(text):
    """Фильтрует речь ведущего, определяя его по количеству предложений"""
    lines = text.split('\n')
    speaker_counts = {}

    # Подсчитываем количество предложений для каждого участника
    for line in lines:
        if ':' in line:
            speaker = line.split(':')[0].strip()
            if speaker and len(speaker) < 50:  # Исключаем слишком длинные имена
                speaker_counts[speaker] = speaker_counts.get(speaker, 0) + 1

    # Находим ведущего (того, кто говорит больше всего)
    if speaker_counts:
        facilitator = max(speaker_counts, key=speaker_counts.get)
        print(f"DEBUG: Определен ведущий: {facilitator} ({speaker_counts[facilitator]} предложений)")

        # Оставляем только речь ведущего
        filtered_lines = []
        for line in lines:
            if line.startswith(f"{facilitator}:"):
                filtered_lines.append(line)

        filtered_text = '\n'.join(filtered_lines)
        print(f"DEBUG: Оригинальный размер: {len(text)} символов")
        print(f"DEBUG: Размер после фильтрации: {len(filtered_text)} символов")
        return filtered_text

    return text


def read_docx_filtered(file_path):
    """Чтение и фильтрация текста из файла .docx"""
    try:
        from docx import Document
        doc = Document(file_path)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)

        original_text = '\n'.join(full_text)
        print(f"DEBUG: Оригинальный размер текста: {len(original_text)} символов")

        # Фильтруем речь ведущего
        filtered_text = filter_facilitator_speech(original_text)
        print(f"DEBUG: Размер после фильтрации: {len(filtered_text)} символов")

        return filtered_text
    except Exception as e:
        error_msg = f"Ошибка чтения файла: {str(e)}"
        print(escape_telegram_chars(error_msg))
        return ""


def save_filtered_text_to_file(filtered_text, filename_prefix="filtered_meeting"):
    """Сохраняет отфильтрованный текст в файл"""
    try:
        # Создаем папку для сохранения отфильтрованных текстов
        os.makedirs("filtered_texts", exist_ok=True)

        # Генерируем имя файла с временной меткой
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"filtered_texts/{filename_prefix}_{timestamp}.txt"

        # Сохраняем текст в файл
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(filtered_text)

        print(f"DEBUG: Отфильтрованный текст сохранен в файл: {filename}")
        return filename
    except Exception as e:
        print(f"Ошибка при сохранении отфильтрованного текста: {e}")
        return None


@dp.message(Command("start"))
async def cmd_start(message: types.Message):
    """Обработчик команды /start"""
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Проверить ММ на соответствие сценарию", callback_data="check_scenario")],
        [InlineKeyboardButton(text="✅ Проверить ММ на соответствие компетенциям модератора ММ",
                              callback_data="check_success_criteria")],
        [InlineKeyboardButton(text="ℹ️ Помощь", callback_data="help")]
    ])

    await message.answer(
        "👋 Добро пожаловать в бот анализа встреч!\n\n"
        "🤖 Я помогу вам проверить встречи на соответствие сценарию и компетенциям модератора.\n\n"
        "Выберите действие:",
        reply_markup=keyboard
    )


@dp.callback_query(lambda c: c.data == "check_success_criteria")
async def check_success_criteria_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик кнопки 'Общие критерии успеха ММ'"""
    await state.set_state(UserStates.waiting_for_success_criteria)

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="❌ Отменить", callback_data="cancel_operation")]
    ])

    await callback.message.answer(
        "🏆 **Проверка компетенций модератора ММ**\n\n"
        "📄 Отправьте файл с текстом встречи (trans.docx) для проверки соответствия копметенциям модератора ММ.\n\n"
        "Для отмены нажмите кнопку ниже или отправьте /cancel",
        reply_markup=keyboard,
        parse_mode="Markdown"
    )
    await callback.answer()


@dp.message(UserStates.waiting_for_success_criteria)
async def handle_success_criteria_file(message: types.Message, state: FSMContext):
    """Обработчик файла для проверки критериев успеха"""
    user_id = message.from_user.id

    if not message.document:
        await message.answer("❌ Пожалуйста, отправьте файл с текстом встречи.")
        return

    file_name = message.document.file_name
    if not file_name.lower().endswith('.docx'):
        await message.answer("❌ Пожалуйста, отправьте файл в формате .docx")
        return

    await message.answer("📥 Загружаю файл для проверки компетенций модератора ММ...")

    try:
        # Скачиваем файл
        file_info = await bot.get_file(message.document.file_id)
        downloaded_file = await bot.download_file(file_info.file_path)

        # Создаем папку для пользователя
        user_folder = f"temp_files/{user_id}"
        os.makedirs(user_folder, exist_ok=True)

        # Сохраняем файл
        file_path = f"{user_folder}/success_criteria_{datetime.now().strftime('%H%M%S')}.docx"
        with open(file_path, 'wb') as f:
            f.write(downloaded_file.read())

        await message.answer("✅ Файл загружен! Проверяю встречу на соответствие компетенциям модератора ММ...")

        # Показываем индикатор обработки
        processing_msg = await message.answer("🔄 Анализирую текст встречи...")

        # Вызываем функцию проверки критериев успеха
        success_result = await check_success_criteria(file_path)

        # Удаляем индикатор обработки
        await bot.delete_message(chat_id=message.chat.id, message_id=processing_msg.message_id)

        # Безопасная отправка результата пользователю
        message_parts = safe_telegram_message(success_result)

        for i, part in enumerate(message_parts):
            if i == 0:
                # Первую часть отправляем с заголовком
                await message.answer(f"🏆 **Результат проверки компетенций модератора ММ:**\n\n{part}",
                                     parse_mode="Markdown")
            else:
                # Остальные части просто как текст
                await message.answer(part)

        # Отправляем кнопку меню
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_to_main"),
             InlineKeyboardButton(text="🔄 Проверить другой файл", callback_data="check_success_criteria")]
        ])

        await message.answer(
            "✅ Проверка завершена!\n\n"
            "Вы можете:\n"
            "• Вернуться в главное меню\n"
            "• Проверить другой файл на компетенции модератора",
            reply_markup=keyboard
        )

        # Очищаем временные файлы
        try:
            os.remove(file_path)
            # Удаляем всю папку пользователя если она пустая
            if os.path.exists(user_folder) and not os.listdir(user_folder):
                os.rmdir(user_folder)
        except Exception as e:
            logging.warning(f"Не удалось очистить временные файлы: {e}")

        await state.clear()

    except Exception as e:
        error_msg = f"❌ Ошибка при проверке компетенций модератора: {str(e)}"
        logging.error(f"Ошибка в handle_success_criteria_file: {e}")

        # Отправляем сообщение об ошибке
        await message.answer(error_msg)

        # Предлагаем вернуться в меню
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_to_main")]
        ])
        await message.answer("Произошла ошибка. Вернитесь в главное меню:", reply_markup=keyboard)

        await state.clear()


@dp.callback_query(lambda c: c.data == "check_scenario")
async def check_scenario_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик кнопки 'Проверить ММ на соответствие сценарию'"""
    # Показываем меню выбора типа ММ
    scenario_keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🏨 ММ онлайн торговля", callback_data="scenario_online")],
        [InlineKeyboardButton(text="🤝 ММ первые встречи", callback_data="scenario_first_meetings")],
        [InlineKeyboardButton(text="📅 ММ первый месяц", callback_data="scenario_first_month")],
        [InlineKeyboardButton(text="🤝 ММ мои встречи", callback_data="scenario_my_meetings")],
        [InlineKeyboardButton(text="🔄 ММ универсальный сценарий", callback_data="scenario_universal")],
        [InlineKeyboardButton(text="❌ Отменить", callback_data="cancel_operation")]
    ])

    await callback.message.answer(
        "📋 Выберите тип ММ для проверки соответствия сценарию:\n\n"
        "🏨 **ММ онлайн торговля** - проверка соответствия сценарию ММ по онлайн торговле\n"
        "🤝 **ММ первые встречи** - проверка первых встреч с клиентами\n"
        "📅 **ММ первый месяц** - проверка ММ за первый месяц работы\n"
        "🤝 **ММ мои встречи** - проверка встреч с клиентами опытных сотрудников\n"
        "🔄 **ММ универсальный сценарий** - универсальная проверка любого типа встреч\n\n"
        "Выберите подходящий вариант:",
        reply_markup=scenario_keyboard,
        parse_mode="Markdown"
    )
    await callback.answer()


@dp.callback_query(lambda c: c.data.startswith("scenario_"))
async def scenario_type_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик выбора типа ММ для проверки сценария"""
    scenario_type = callback.data

    # Сохраняем выбранный тип сценария в состоянии
    await state.update_data(scenario_type=scenario_type)
    await state.set_state(UserStates.waiting_for_scenario_check)

    # Определяем название типа для отображения
    scenario_names = {
        "scenario_online": "ММ онлайн торговля",
        "scenario_first_meetings": "ММ первые встречи",
        "scenario_first_month": "ММ первый месяц",
        "scenario_my_meetings": "ММ мои встречи",
        "scenario_universal": "ММ универсальный сценарий"
    }

    scenario_name = scenario_names.get(scenario_type, "выбранный тип ММ")

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="❌ Отменить", callback_data="cancel_operation")]
    ])

    await callback.message.answer(
        f"📄 Для проверки *{scenario_name}* на соответствие сценарию отправьте файл с текстом встречи (trans.docx)\n\n"
        "Для отмена нажмите кнопку ниже или отправьте /cancel",
        reply_markup=keyboard,
        parse_mode="Markdown"
    )
    await callback.answer()


@dp.message(UserStates.waiting_for_scenario_check)
async def handle_scenario_check_file(message: types.Message, state: FSMContext):
    """Обработчик файла для проверки сценария"""
    user_id = message.from_user.id

    if not message.document:
        await message.answer("❌ Пожалуйста, отправьте файл с текстом встречи.")
        return

    file_name = message.document.file_name
    if not file_name.lower().endswith('.docx'):
        await message.answer("❌ Пожалуйста, отправьте файл в формате .docx")
        return

    # Получаем выбранный тип сценария из состояния
    user_data = await state.get_data()
    scenario_type = user_data.get('scenario_type', 'scenario_online')

    scenario_names = {
        "scenario_online": "ММ онлайн торговля",
        "scenario_first_meetings": "ММ первые встречи",
        "scenario_first_month": "ММ первый месяц",
        "scenario_my_meetings": "ММ мои встречи",
        "scenario_universal": "ММ универсальный сценарий"
    }

    scenario_name = scenario_names.get(scenario_type, "ММ")

    await message.answer(f"📥 Загружаю файл для проверки {scenario_name}...")

    try:
        # Скачиваем файл
        file_info = await bot.get_file(message.document.file_id)
        downloaded_file = await bot.download_file(file_info.file_path)

        # Создаем папку для пользователя
        user_folder = f"temp_files/{user_id}"
        os.makedirs(user_folder, exist_ok=True)

        # Сохраняем файл
        file_path = f"{user_folder}/scenario_check_{datetime.now().strftime('%H%M%S')}.docx"
        with open(file_path, 'wb') as f:
            f.write(downloaded_file.read())

        # Читаем и фильтруем текст
        print("\n" + "=" * 80)
        print("DEBUG: ОБРАБОТКА ТЕКСТА ВСТРЕЧИ")
        print("=" * 80)

        # Читаем оригинальный текст
        from docx import Document
        doc = Document(file_path)
        original_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])

        print(f"ОРИГИНАЛЬНЫЙ ТЕКСТ ({len(original_text)} символов):")
        print("-" * 40)
        print(original_text[:1000] + "..." if len(original_text) > 1000 else original_text)
        print("-" * 40)

        # Фильтруем текст (оставляем только речь ведущего)
        filtered_text = read_docx_filtered(file_path)

        print(f"ОТФИЛЬТРОВАННЫЙ ТЕКСТ ({len(filtered_text)} символов):")
        print("-" * 40)
        print(filtered_text[:1000] + "..." if len(filtered_text) > 1000 else filtered_text)
        print("-" * 40)
        print(f"Сокращение текста: {len(original_text)} -> {len(filtered_text)} символов "
              f"({(1 - len(filtered_text) / len(original_text)) * 100:.1f}% сокращено)")

        # СОХРАНЯЕМ ОТФИЛЬТРОВАННЫЙ ТЕКСТ В ФАЙЛ
        saved_file_path = save_filtered_text_to_file(filtered_text, f"scenario_{scenario_type}")
        if saved_file_path:
            print(f"DEBUG: Весь отфильтрованный текст сохранен в: {saved_file_path}")

        print("=" * 80 + "\n")

        await message.answer(f"✅ Файл загружен! Проверяю {scenario_name} на соответствие сценарию...")

        # Показываем индикатор обработки
        processing_msg = await message.answer("🔄 Анализирую текст встречи...")

        # СОЗДАЕМ ВРЕМЕННЫЙ ФАЙЛ С ОТФИЛЬТРОВАННЫМ ТЕКСТОМ ДЛЯ ОТПРАВКИ В GIGACHAT
        filtered_file_path = f"{user_folder}/filtered_for_gigachat.docx"

        # Создаем новый DOCX файл с отфильтрованным текстом
        from docx import Document as DocxDocument
        filtered_doc = DocxDocument()
        for line in filtered_text.split('\n'):
            if line.strip():
                filtered_doc.add_paragraph(line)
        filtered_doc.save(filtered_file_path)

        # Вызываем функцию проверки сценария с ОТФИЛЬТРОВАННЫМ файлом
        scenario_result = await check_meeting_scenario(filtered_file_path, scenario_type)

        # Удаляем индикатор обработки
        await bot.delete_message(chat_id=message.chat.id, message_id=processing_msg.message_id)

        # Безопасная отправка результата пользователю
        message_parts = safe_telegram_message(scenario_result)

        for i, part in enumerate(message_parts):
            if i == 0:
                # Первую часть отправляем с заголовком
                await message.answer(part, parse_mode="Markdown")
            else:
                # Остальные части просто как текст
                await message.answer(part)

        # Отправляем кнопку меню
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_to_main"),
             InlineKeyboardButton(text="🔄 Проверить другой сценарий", callback_data="check_scenario")]
        ])

        await message.answer(
            "✅ Проверка завершена!\n\n"
            "Вы можете:\n"
            "• Вернуться в главное меню\n"
            "• Проверить другой сценарий",
            reply_markup=keyboard
        )

        # Очищаем временные файлы
        try:
            os.remove(file_path)
            os.remove(filtered_file_path)
            # Удаляем всю папку пользователя если она пустая
            if os.path.exists(user_folder) and not os.listdir(user_folder):
                os.rmdir(user_folder)
        except Exception as e:
            logging.warning(f"Не удалось очистить временные файлы: {e}")

        await state.clear()

    except Exception as e:
        error_msg = f"❌ Ошибка при проверке сценария: {str(e)}"
        logging.error(f"Ошибка в handle_scenario_check_file: {e}")

        # Отправляем сообщение об ошибке
        await message.answer(error_msg)

        # Предлагаем вернуться в меню
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_to_main")]
        ])
        await message.answer("Произошла ошибка. Вернитесь в главное меню:", reply_markup=keyboard)

        await state.clear()


@dp.callback_query(lambda c: c.data == "help")
async def help_handler(callback: types.CallbackQuery):
    """Обработчик кнопки 'Помощь'"""
    help_text = """
ℹ️ **ПОМОЩЬ ПО ИСПОЛЬЗОВАНИЮ БОТА**

✅ **Проверить ММ на соответствие сценарию:**
• Выберите тип ММ: онлайн торговля, первые встречи, первый месяц, мои встречи или универсальный сценарий
• Загрузите файл с текстом встречи (trans.docx)
• Получите анализ соответствия проведенной встречи запланированному сценарию

🏆 **Проверить на соответствие компетенциям модератора ММ:**
• Загрузите файл с текстом встречи (trans.docx)
• Получите проверку соответствия копметенциям модератора ММ

📁 **Форматы файлов:**
• trans.docx - текст встречи с временными метками

⏱️ **Время обработки:**
• Проверка сценария: 10-20 секунд
• Проверка компетенций: 10-20 секунд

🔧 **При возникновении проблем:**
• Проверьте формат файлов
• Убедитесь в наличии интернета
• Попробуйте загрузить файлы заново
"""

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🔙 Назад", callback_data="back_to_main")]
    ])

    await callback.message.edit_text(help_text, reply_markup=keyboard, parse_mode="Markdown")
    await callback.answer()


@dp.callback_query(lambda c: c.data == "back_to_main")
async def back_to_main_handler(callback: types.CallbackQuery):
    """Обработчик кнопки 'Назад'"""
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Проверить ММ на соответствие сценарию", callback_data="check_scenario")],
        [InlineKeyboardButton(text="✅ Проверить на соответствие компетенциям модератора ММ",
                              callback_data="check_success_criteria")],
        [InlineKeyboardButton(text="ℹ️ Помощь", callback_data="help")]
    ])

    await callback.message.edit_text(
        "👋 Добро пожаловать в бот анализа встреч!\n\n"
        "🤖 Я помогу вам проверить встречи на соответствие сценарию и компетенциям модератора.\n\n"
        "Выберите действие:",
        reply_markup=keyboard
    )
    await callback.answer()


@dp.callback_query(lambda c: c.data == "cancel_operation")
async def cancel_operation_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик отмены операции с очисткой временных файлов"""
    user_id = callback.from_user.id
    user_folder = f"temp_files/{user_id}"

    # Очищаем временные файлы
    try:
        if os.path.exists(user_folder):
            import shutil
            shutil.rmtree(user_folder)
    except Exception as e:
        logging.warning(f"Не удалось очистить временные файлы при отмене: {e}")

    await state.clear()

    # Возвращаем в главное меню
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Проверить ММ на соответствие сценарию", callback_data="check_scenario")],
        [InlineKeyboardButton(text="✅ Проверить на соответствие компетенциям модератора ММ",
                              callback_data="check_success_criteria")],
        [InlineKeyboardButton(text="ℹ️ Помощь", callback_data="help")]
    ])

    await callback.message.edit_text(
        "❌ Операция отменена. Выберите новое действие:",
        reply_markup=keyboard
    )
    await callback.answer()


async def main():
    """Главная функция"""
    # Создаем папки для временных файлов
    os.makedirs("temp_files", exist_ok=True)
    os.makedirs("filtered_texts", exist_ok=True)

    # Запускаем бота
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())