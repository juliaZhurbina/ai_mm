import asyncio
import logging
from aiogram import Bot, Dispatcher, types
from aiogram.filters import Command
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
import os
from datetime import datetime
from scenario_checker import check_meeting_scenario, safe_telegram_message

# Импортируем наши модули
import sys

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from giga_recomendation import MeetingAnalyzer
from competency_analyzer import analyze_competencies_async
from docx import Document

# Импортируем модуль для проверки критериев успеха
from success_criteria import check_success_criteria

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)

# Конфигурация бота
BOT_TOKEN = "8079592721:AAGLaX7LwUPX0X5fr1SK-9IQnSvfP3Z96ws"
bot = Bot(token=BOT_TOKEN)
storage = MemoryStorage()
dp = Dispatcher(storage=storage)


# Состояния FSM
class UserStates(StatesGroup):
    waiting_for_files = State()
    waiting_for_trans_file = State()
    waiting_for_triggers_file = State()
    waiting_for_scenario_check = State()
    waiting_for_success_criteria = State()  # Новое состояние для проверки критериев успеха


# Создаем экземпляр анализатора встреч
AUTH_KEY = 'ZGMzMGJmZjEtODQwYS00ZjAwLWI2NjgtNGIyNGNiY2ViNmE1OjYwNjM3NTU0LWQxMDctNDA5ZS1hZWM3LTAwYjQ5MjZkOGU2OA=='
SCOPE = 'GIGACHAT_API_PERS'
API_AUTH_URL = 'https://ngw.devices.sberbank.ru:9443/api/v2/oauth'
API_CHAT_URL = 'https://gigachat.devices.sberbank.ru/api/v1/chat/completions'

meeting_analyzer = MeetingAnalyzer(AUTH_KEY, SCOPE, API_AUTH_URL, API_CHAT_URL)

# Словарь для хранения файлов пользователей
user_files = {}


def escape_telegram_chars(text: str) -> str:
    """Экранирование специальных символов для Telegram"""
    escape_chars = ['_', '*', '[', ']', '(', ')', '~', '`', '>', '#', '+', '-', '=', '|', '{', '}', '.', '!']
    for char in escape_chars:
        text = text.replace(char, f'\\{char}')
    return text


def filter_facilitator_speech(text):
    """Фильтрует речь ведущего, определяя его по количеству предложений"""
    lines = text.split('\n')
    speaker_counts = {}

    # Подсчитываем количество предложений для каждого участника
    for line in lines:
        if ':' in line:
            speaker = line.split(':')[0].strip()
            if speaker and len(speaker) < 50:  # Исключаем слишком длинные имена
                speaker_counts[speaker] = speaker_counts.get(speaker, 0) + 1

    # Находим ведущего (того, кто говорит больше всего)
    if speaker_counts:
        facilitator = max(speaker_counts, key=speaker_counts.get)
        print(f"DEBUG: Определен ведущий: {facilitator} ({speaker_counts[facilitator]} предложений)")

        # Оставляем только речь ведущего
        filtered_lines = []
        for line in lines:
            if line.startswith(f"{facilitator}:"):
                filtered_lines.append(line)

        filtered_text = '\n'.join(filtered_lines)
        print(f"DEBUG: Оригинальный размер: {len(text)} символов")
        print(f"DEBUG: Размер после фильтрации: {len(filtered_text)} символов")
        return filtered_text

    return text


def read_docx_filtered(file_path):
    """Чтение и фильтрация текста из файла .docx"""
    try:
        doc = Document(file_path)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)

        original_text = '\n'.join(full_text)
        print(f"DEBUG: Оригинальный размер текста: {len(original_text)} символов")

        # Фильтруем речь ведущего
        filtered_text = filter_facilitator_speech(original_text)
        print(f"DEBUG: Размер после фильтрации: {len(filtered_text)} символов")

        return filtered_text
    except Exception as e:
        error_msg = f"Ошибка чтения файла: {str(e)}"
        print(escape_telegram_chars(error_msg))
        return ""


@dp.message(Command("start"))
async def cmd_start(message: types.Message):
    """Обработчик команды /start"""
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📊 Анализ компетенций", callback_data="analyze_competencies")],
        [InlineKeyboardButton(text="💡 Рекомендации", callback_data="get_recommendations")],
        [InlineKeyboardButton(text="✅ Проверить ММ на соответствие сценарию", callback_data="check_scenario")],
        [InlineKeyboardButton(text="🏆 Общие критерии успеха ММ", callback_data="check_success_criteria")],
        # Новая кнопка
        [InlineKeyboardButton(text="ℹ️ Помощь", callback_data="help")]
    ])

    await message.answer(
        "👋 Добро пожаловать в бот анализа встреч!\n\n"
        "🤖 Я помогу вам проанализировать компетенции и получить рекомендации по развитию.\n\n"
        "Выберите действие:",
        reply_markup=keyboard
    )


@dp.callback_query(lambda c: c.data == "check_success_criteria")
async def check_success_criteria_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик кнопки 'Общие критерии успеха ММ'"""
    await state.set_state(UserStates.waiting_for_success_criteria)

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="❌ Отменить", callback_data="cancel_operation")]
    ])

    await callback.message.answer(
        "🏆 **Проверка общих критериев успеха ММ**\n\n"
        "📄 Отправьте файл с текстом встречи (trans.docx) для проверки по общим критериям успеха ММ.\n\n"
        "Критерии включают:\n"
        "• Структура встречи\n"
        "• Управление временем\n"
        "• Вовлечение участников\n"
        "• Достижение целей встречи\n"
        "• И другие важные аспекты\n\n"
        "Для отмены нажмите кнопку ниже или отправьте /cancel",
        reply_markup=keyboard,
        parse_mode="Markdown"
    )
    await callback.answer()

@dp.message(UserStates.waiting_for_success_criteria)
async def handle_success_criteria_file(message: types.Message, state: FSMContext):
    """Обработчик файла для проверки критериев успеха"""
    user_id = message.from_user.id

    if not message.document:
        await message.answer("❌ Пожалуйста, отправьте файл с текстом встречи.")
        return

    file_name = message.document.file_name
    if not file_name.lower().endswith('.docx'):
        await message.answer("❌ Пожалуйста, отправьте файл в формате .docx")
        return

    await message.answer("📥 Загружаю файл для проверки критериев успеха ММ...")

    try:
        # Скачиваем файл
        file_info = await bot.get_file(message.document.file_id)
        downloaded_file = await bot.download_file(file_info.file_path)

        # Создаем папку для пользователя
        user_folder = f"temp_files/{user_id}"
        os.makedirs(user_folder, exist_ok=True)

        # Сохраняем файл
        file_path = f"{user_folder}/success_criteria_{datetime.now().strftime('%H%M%S')}.docx"
        with open(file_path, 'wb') as f:
            f.write(downloaded_file.read())

        await message.answer("✅ Файл загружен! Проверяю встречу по общим критериям успеха ММ...")

        # Показываем индикатор обработки
        processing_msg = await message.answer("🔄 Анализирую текст встречи...")

        # Вызываем функцию проверки критериев успеха
        success_result = await check_success_criteria(file_path)

        # Удаляем индикатор обработки
        await bot.delete_message(chat_id=message.chat.id, message_id=processing_msg.message_id)

        # Безопасная отправка результата пользователю
        message_parts = safe_telegram_message(success_result)

        for i, part in enumerate(message_parts):
            if i == 0:
                # Первую часть отправляем с заголовком
                await message.answer(f"🏆 **Результат проверки критериев успеха ММ:**\n\n{part}",
                                     parse_mode="Markdown")
            else:
                # Остальные части просто как текст
                await message.answer(part)

        # Отправляем кнопку меню
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_to_main"),
             InlineKeyboardButton(text="🔄 Проверить другой файл", callback_data="check_success_criteria")]
        ])

        await message.answer(
            "✅ Проверка завершена!\n\n"
            "Вы можете:\n"
            "• Вернуться в главное меню\n"
            "• Проверить другой файл по критериям успеха",
            reply_markup=keyboard
        )

        # Очищаем временные файлы
        try:
            os.remove(file_path)
            # Удаляем всю папку пользователя если она пустая
            if os.path.exists(user_folder) and not os.listdir(user_folder):
                os.rmdir(user_folder)
        except Exception as e:
            logging.warning(f"Не удалось очистить временные файлы: {e}")

        await state.clear()

    except Exception as e:
        error_msg = f"❌ Ошибка при проверке критериев успеха: {str(e)}"
        logging.error(f"Ошибка в handle_success_criteria_file: {e}")

        # Отправляем сообщение об ошибке
        await message.answer(error_msg)

        # Предлагаем вернуться в меню
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_to_main")]
        ])
        await message.answer("Произошла ошибка. Вернитесь в главное меню:", reply_markup=keyboard)

        await state.clear()


@dp.callback_query(lambda c: c.data == "check_scenario")
async def check_scenario_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик кнопки 'Проверить ММ на соответствие сценарию'"""
    # Показываем меню выбора типа ММ
    scenario_keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🏨 ММ онлайн торговля", callback_data="scenario_online")],
        [InlineKeyboardButton(text="🤝 ММ первые встречи", callback_data="scenario_first_meetings")],
        [InlineKeyboardButton(text="📅 ММ первый месяц", callback_data="scenario_first_month")],
        [InlineKeyboardButton(text="❌ Отменить", callback_data="cancel_operation")]
    ])

    await callback.message.answer(
        "📋 Выберите тип ММ для проверки соответствия сценарию:\n\n"
        "🏨 **ММ онлайн торговля** - проверка соответствия сценарию ММ по онлайн торговле\n"
        "🤝 **ММ первые встречи** - проверка первых встреч с клиентами\n"
        "📅 **ММ первый месяц** - проверка ММ за первый месяц работы\n\n"
        "Выберите подходящий вариант:",
        reply_markup=scenario_keyboard,
        parse_mode="Markdown"
    )
    await callback.answer()


@dp.callback_query(lambda c: c.data.startswith("scenario_"))
async def scenario_type_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик выбора типа ММ для проверки сценария"""
    scenario_type = callback.data

    # Сохраняем выбранный тип сценария в состоянии
    await state.update_data(scenario_type=scenario_type)
    await state.set_state(UserStates.waiting_for_scenario_check)

    # Определяем название типа для отображения
    scenario_names = {
        "scenario_online": "ММ онлайн торговля",
        "scenario_first_meetings": "ММ первые встречи",
        "scenario_first_month": "ММ первый месяц"
    }

    scenario_name = scenario_names.get(scenario_type, "выбранный тип ММ")

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="❌ Отменить", callback_data="cancel_operation")]
    ])

    await callback.message.answer(
        f"📄 Для проверки *{scenario_name}* на соответствие сценарию отправьте файл с текстом встречи (trans.docx)\n\n"
        "Для отмены нажмите кнопку ниже или отправьте /cancel",
        reply_markup=keyboard,
        parse_mode="Markdown"
    )
    await callback.answer()


@dp.message(UserStates.waiting_for_scenario_check)
async def handle_scenario_check_file(message: types.Message, state: FSMContext):
    """Обработчик файла для проверки сценария"""
    user_id = message.from_user.id

    if not message.document:
        await message.answer("❌ Пожалуйста, отправьте файл с текстом встречи.")
        return

    file_name = message.document.file_name
    if not file_name.lower().endswith('.docx'):
        await message.answer("❌ Пожалуйста, отправьте файл в формате .docx")
        return

    # Получаем выбранный тип сценария из состояния
    user_data = await state.get_data()
    scenario_type = user_data.get('scenario_type', 'scenario_online')

    scenario_names = {
        "scenario_online": "ММ онлайн торговля",
        "scenario_first_meetings": "ММ первые встречи",
        "scenario_first_month": "ММ первый месяц"
    }

    scenario_name = scenario_names.get(scenario_type, "ММ")

    await message.answer(f"📥 Загружаю файл для проверки {scenario_name}...")

    try:
        # Скачиваем файл
        file_info = await bot.get_file(message.document.file_id)
        downloaded_file = await bot.download_file(file_info.file_path)

        # Создаем папку для пользователя
        user_folder = f"temp_files/{user_id}"
        os.makedirs(user_folder, exist_ok=True)

        # Сохраняем файл
        file_path = f"{user_folder}/scenario_check_{datetime.now().strftime('%H%M%S')}.docx"
        with open(file_path, 'wb') as f:
            f.write(downloaded_file.read())

        await message.answer(f"✅ Файл загружен! Проверяю {scenario_name} на соответствие сценарию...")

        # Показываем индикатор обработки
        processing_msg = await message.answer("🔄 Анализирую текст встречи...")

        # Вызываем функцию проверки сценария
        scenario_result = await check_meeting_scenario(file_path, scenario_type)

        # Удаляем индикатор обработки
        await bot.delete_message(chat_id=message.chat.id, message_id=processing_msg.message_id)

        # Безопасная отправка результата пользователю
        message_parts = safe_telegram_message(scenario_result)

        for i, part in enumerate(message_parts):
            if i == 0:
                # Первую часть отправляем с заголовком
                await message.answer(part, parse_mode="Markdown")
            else:
                # Остальные части просто как текст
                await message.answer(part)

        # Отправляем кнопку меню
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_to_main"),
             InlineKeyboardButton(text="🔄 Проверить другой сценарий", callback_data="check_scenario")]
        ])

        await message.answer(
            "✅ Проверка завершена!\n\n"
            "Вы можете:\n"
            "• Вернуться в главное меню\n"
            "• Проверить другой сценарий",
            reply_markup=keyboard
        )

        # Очищаем временные файлы
        try:
            os.remove(file_path)
            # Удаляем всю папку пользователя если она пустая
            if os.path.exists(user_folder) and not os.listdir(user_folder):
                os.rmdir(user_folder)
        except Exception as e:
            logging.warning(f"Не удалось очистить временные файлы: {e}")

        await state.clear()

    except Exception as e:
        error_msg = f"❌ Ошибка при проверке сценария: {str(e)}"
        logging.error(f"Ошибка в handle_scenario_check_file: {e}")

        # Отправляем сообщение об ошибке
        await message.answer(error_msg)

        # Предлагаем вернуться в меню
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_to_main")]
        ])
        await message.answer("Произошла ошибка. Вернитесь в главное меню:", reply_markup=keyboard)

        await state.clear()

@dp.callback_query(lambda c: c.data == "analyze_competencies")
async def analyze_competencies_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик кнопки 'Анализ компетенций'"""
    await state.set_state(UserStates.waiting_for_files)
    user_id = callback.from_user.id
    user_folder = f"temp_files/{user_id}"
    if os.path.exists(user_folder):
        import shutil
        shutil.rmtree(user_folder)
    os.makedirs(user_folder, exist_ok=True)
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="❌ Отменить", callback_data="cancel_operation")]
    ])
    await callback.message.answer(
        "📁 Пожалуйста, загрузите оба файла:\n"
        "1. Файл с текстом встречи (trans.docx)\n"
        "2. Файл с триггерами (triggers.xlsx)\n\n"
        "Отправьте файлы по одному.\n\n"
        "💡 Для отмены нажмите кнопку ниже или отправьте /cancel",
        reply_markup=keyboard
    )
    await callback.answer()


@dp.callback_query(lambda c: c.data == "get_recommendations")
async def get_recommendations_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик кнопки 'Рекомендации'"""
    if not os.path.exists('REPORT.txt'):
        await callback.message.answer(
            "❌ Сначала выполните анализ компетенций!\n\n"
            "📊 Нажмите кнопку 'Анализ компетенций' и загрузите оба файла:\n"
            "• trans.docx (текст встречи)\n"
            "• triggers.xlsx (триггеры компетенций)\n\n"
            "После получения отчета по компетенциям сможете получить детальные рекомендации."
        )
        await callback.answer()
        return

    user_id = callback.from_user.id
    user_folder = f"temp_files/{user_id}"
    trans_file_path = f"{user_folder}/trans.docx"
    if os.path.exists(trans_file_path):
        await callback.message.answer("✅ Начинаю анализ...")
        try:
            analysis_result = meeting_analyzer.analyze_meeting_with_file(trans_file_path)
            if analysis_result.startswith("❌"):
                await callback.message.answer(analysis_result)
                os.remove(trans_file_path)
                await state.clear()
                return

            # Создаем настоящий DOCX файл
            recommendations_filename = f"detailed_recommendations_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

            doc = Document()
            doc.add_heading('Детальные рекомендации по развитию компетенций', 0)

            paragraphs = analysis_result.split('\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph)

            doc.save(recommendations_filename)

            with open(recommendations_filename, 'rb') as f:
                await callback.message.answer_document(
                    types.BufferedInputFile(f.read(), filename=recommendations_filename),
                    caption="💡 Детальные рекомендации по развитию компетенций\n\n📊 Анализ основан на:\n• Тексте встречи\n• Отчете по компетенциям"
                )

            # Удаляем временный файл
            os.remove(recommendations_filename)

            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="🏠 Меню", callback_data="back_to_main")]
            ])
            await callback.message.answer("✅ Анализ завершен! Нажмите кнопку ниже для возврата в главное меню:",
                                          reply_markup=keyboard)
            os.remove(trans_file_path)
            await state.clear()
        except Exception as e:
            error_msg = f"❌ Ошибка при анализе: {str(e)}"
            safe_error = safe_telegram_message(error_msg)[0]
            await callback.message.answer(safe_error, parse_mode=None)
            os.remove(trans_file_path)
            await state.clear()
        return

    await state.set_state(UserStates.waiting_for_trans_file)
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="❌ Отменить", callback_data="cancel_operation")]
    ])
    await callback.message.answer(
        "📄 Пожалуйста, загрузите файл с текстом встречи (trans.docx)\n\n"
        "💡 Для отмены нажмите кнопку ниже или отправьте /cancel",
        reply_markup=keyboard
    )
    await callback.answer()


@dp.callback_query(lambda c: c.data == "help")
async def help_handler(callback: types.CallbackQuery):
    """Обработчик кнопки 'Помощь'"""
    help_text = """
ℹ️ **ПОМОЩЬ ПО ИСПОЛЬЗОВАНИЮ БОТА**

📊 **Анализ компетенций:**
• Загрузите файл с текстом встречи (trans.docx)
• Загрузите файл с триггерами (triggers.xlsx)
• Получите полный отчет и краткое резюме
• Узнайте баллы по компетенции и рекомендуемые курсы

💡 **Рекомендации:**
• Сначала выполните анализ компетенций
• Затем загрузите файл с текстом встречи (trans.docx)
• Получите детальные рекомендации на основе анализа встречи и компетенций
• Конкретные курсы и план развития на месяц

✅ **Проверить ММ на соответствие сценарию:**
• Выберите тип ММ: гостиницы, первые встречи или первый месяц
• Загрузите файл с текстом встречи (trans.docx)
• Получите анализ соответствия проведенной встречи запланированному сценарию

📁 **Форматы файлов:**
• trans.docx - текст встречи с временными метками
• triggers.xlsx - Excel файл с триггерами компетенций

⏱️ **Время обработки:**
• Анализ компетенций: 20-30 секунд
• Рекомендации: 10-15 секунд
• Проверка сценария: 10-20 секунд

🔧 **При возникновении проблем:**
• Проверьте формат файлов
• Убедитесь в наличии интернета
• Попробуйте загрузить файлы заново
"""

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🔙 Назад", callback_data="back_to_main")]
    ])

    await callback.message.edit_text(help_text, reply_markup=keyboard, parse_mode="Markdown")
    await callback.answer()


@dp.callback_query(lambda c: c.data == "back_to_main")
async def back_to_main_handler(callback: types.CallbackQuery):
    """Обработчик кнопки 'Назад'"""
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📊 Анализ компетенций", callback_data="analyze_competencies")],
        [InlineKeyboardButton(text="💡 Рекомендации", callback_data="get_recommendations")],
        [InlineKeyboardButton(text="✅ Проверить ММ на соответствие сценарию", callback_data="check_scenario")],
        [InlineKeyboardButton(text="🏆 Общие критерии успеха ММ", callback_data="check_success_criteria")],  # Новая кнопка
        [InlineKeyboardButton(text="ℹ️ Помощь", callback_data="help")]
    ])

    await callback.message.edit_text(
        "👋 Добро пожаловать в бот анализа встреч!\n\n"
        "🤖 Я помогу вам проанализировать компетенции и получить рекомендации по развитию.\n\n"
        "Выберите действие:",
        reply_markup=keyboard
    )
    await callback.answer()

def get_main_keyboard():
    """Возвращает главную клавиатуру"""
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📊 Анализ компетенций", callback_data="analyze_competencies")],
        [InlineKeyboardButton(text="💡 Рекомендации", callback_data="get_recommendations")],
        [InlineKeyboardButton(text="✅ Проверить ММ на соответствие сценарию", callback_data="check_scenario")],
        [InlineKeyboardButton(text="🏆 Общие критерии успеха ММ", callback_data="check_success_criteria")],  # Новая кнопка
        [InlineKeyboardButton(text="ℹ️ Помощь", callback_data="help")]
    ])

@dp.callback_query(lambda c: c.data == "cancel_operation")
async def cancel_operation_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик отмены операции с очисткой временных файлов"""
    user_id = callback.from_user.id
    user_folder = f"temp_files/{user_id}"

    # Очищаем временные файлы
    try:
        if os.path.exists(user_folder):
            import shutil
            shutil.rmtree(user_folder)
    except Exception as e:
        logging.warning(f"Не удалось очистить временные файлы при отмене: {e}")

    await state.clear()

    # Возвращаем в главное меню
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📊 Анализ компетенций", callback_data="analyze_competencies")],
        [InlineKeyboardButton(text="💡 Рекомендации", callback_data="get_recommendations")],
        [InlineKeyboardButton(text="✅ Проверить ММ на соответствие сценарию", callback_data="check_scenario")],
        [InlineKeyboardButton(text="ℹ️ Помощь", callback_data="help")]
    ])

    await callback.message.edit_text(
        "❌ Операция отменена. Выберите новое действие:",
        reply_markup=keyboard
    )
    await callback.answer()


@dp.message(UserStates.waiting_for_files)
async def handle_files_for_analysis(message: types.Message, state: FSMContext):
    """Обработчик файлов для анализа компетенций"""
    user_id = message.from_user.id

    # ОЧИСТКА temp_files только при новом анализе компетенций
    user_folder = f"temp_files/{user_id}"
    if not os.path.exists(user_folder):
        os.makedirs(user_folder, exist_ok=True)

    if not message.document:
        await message.answer("❌ Пожалуйста, отправьте файл.")
        return

    file_name = message.document.file_name
    # Сохраняем docx как trans.docx, xlsx как triggers.xlsx
    if file_name.lower().endswith('.docx'):
        save_name = 'trans.docx'
    elif file_name.lower().endswith('.xlsx'):
        save_name = 'triggers.xlsx'
    else:
        save_name = file_name  # fallback

    file_path = f"{user_folder}/{save_name}"
    file_info = await bot.get_file(message.document.file_id)
    downloaded_file = await bot.download_file(file_info.file_path)
    with open(file_path, 'wb') as f:
        f.write(downloaded_file.read())

    # Если это docx файл, сразу фильтруем речь ведущего
    if file_name.lower().endswith('.docx'):
        print(f"DEBUG: Фильтрую речь ведущего из файла {file_name}")
        filtered_text = read_docx_filtered(file_path)

        # Сохраняем отфильтрованный текст обратно в файл
        from docx import Document
        doc = Document()
        doc.add_paragraph(filtered_text)
        doc.save(file_path)
        print(f"DEBUG: Сохранен отфильтрованный текст в {file_path}")

    if user_id not in user_files:
        user_files[user_id] = {}

    user_files[user_id][save_name] = file_path

    await message.answer(f"✅ Файл {file_name} загружен!")

    # Проверяем, есть ли оба файла
    if len(user_files[user_id]) == 2:
        await analyze_user_files(message, user_id)
        await state.clear()
    else:
        await message.answer("📁 Теперь отправьте второй файл.")


async def analyze_user_files(message: types.Message, user_id: int):
    """Анализ файлов пользователя"""
    try:
        files = user_files[user_id]

        # Находим файлы
        trans_file = None
        triggers_file = None

        for file_name, file_path in files.items():
            if file_name.endswith('.docx'):
                trans_file = file_path
            elif file_name.endswith('.xlsx'):
                triggers_file = file_path

        if not trans_file or not triggers_file:
            await message.answer("❌ Не удалось найти оба файла. Попробуйте снова.")
            return

        # Анализируем компетенции
        await message.answer("🔍 Анализирую компетенции...")

        # Анализируем компетенции
        full_report, summary = await analyze_competencies_async(trans_file, triggers_file)

        # Сохраняем полный отчет
        report_filename = f"competency_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        with open(report_filename, 'w', encoding='utf-8') as f:
            f.write(full_report)

        # Обновляем REPORT.txt актуальными данными
        with open('REPORT.txt', 'w', encoding='utf-8') as f:
            f.write(full_report)

        # Отправляем файл с отчетом
        with open(report_filename, 'rb') as f:
            await message.answer_document(
                types.BufferedInputFile(f.read(), filename=report_filename),
                caption="📊 Полный отчет по анализу компетенций"
            )

        # Сначала отправляем краткое резюме без кнопки
        safe_summary = safe_telegram_message(summary)[0]
        await message.answer(f"📋 {safe_summary}", parse_mode=None)

        # Затем отдельное сообщение с кнопкой возврата в меню
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Меню", callback_data="back_to_main")]
        ])
        await message.answer("Вы можете вернуться в меню:", reply_markup=keyboard)

    except Exception as e:
        error_message = f"❌ Ошибка при анализе: {str(e)}"
        safe_error = safe_telegram_message(error_message)[0]
        logging.error(f"Ошибка анализа компетенций для пользователя {user_id}: {e}")
        await message.answer(safe_error, parse_mode=None)


@dp.message(UserStates.waiting_for_trans_file)
async def handle_trans_file(message: types.Message, state: FSMContext):
    """Обработчик файла встречи для рекомендаций"""
    user_id = message.from_user.id

    if not message.document:
        await message.answer("❌ Пожалуйста, отправьте файл.")
        return

    file_name = message.document.file_name
    if not file_name.endswith('.docx'):
        await message.answer("❌ Пожалуйста, отправьте файл в формате .docx")
        return

    file_info = await bot.get_file(message.document.file_id)
    downloaded_file = await bot.download_file(file_info.file_path)

    # Создаем папку для пользователя
    user_folder = f"temp_files/{user_id}"
    os.makedirs(user_folder, exist_ok=True)

    # Сохраняем файл
    file_path = f"{user_folder}/{file_name}"
    with open(file_path, 'wb') as f:
        f.write(downloaded_file.read())

    await message.answer("✅ Файл загружен! Анализирую встречу и отчет по компетенциям...")

    try:
        # Анализируем встречу с учетом отчета компетенций
        analysis_result = meeting_analyzer.analyze_meeting_with_file(file_path)

        # Проверяем, не является ли результат ошибкой
        if analysis_result.startswith("❌"):
            await message.answer(analysis_result)
            os.remove(file_path)
            await state.clear()
            return

        # Создаем настоящий DOCX файл с детальными рекомендациями
        recommendations_filename = f"detailed_recommendations_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        # Создаем документ и добавляем текст
        doc = Document()

        # Добавляем заголовок
        doc.add_heading('Детальные рекомендации по развитию компетенций', 0)

        # Разбиваем текст на абзацы и добавляем в документ
        paragraphs = analysis_result.split('\n')
        for paragraph in paragraphs:
            if paragraph.strip():  # Пропускаем пустые строки
                doc.add_paragraph(paragraph)

        # Сохраняем документ
        doc.save(recommendations_filename)

        # Отправляем файл с детальными рекомендациями
        with open(recommendations_filename, 'rb') as f:
            await message.answer_document(
                types.BufferedInputFile(f.read(), filename=recommendations_filename),
                caption="💡 Детальные рекомендации по развитию компетенций\n\n📊 Анализ основан на:\n• Тексте встречи\n• Отчете по компетенциям"
            )

        # Удаляем временный файл
        os.remove(recommendations_filename)

        # Отправляем кнопку меню
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Меню", callback_data="back_to_main")]
        ])
        await message.answer("✅ Анализ завершен! Нажмите кнопку ниже для возврата в главное меню:",
                             reply_markup=keyboard)

        # Очищаем файл
        os.remove(file_path)
        await state.clear()

    except Exception as e:
        await message.answer(f"❌ Ошибка при анализе: {str(e)}")
        os.remove(file_path)
        await state.clear()

def cleanup_user_files(user_id):
    """Очищает файлы пользователя"""
    if user_id in user_files:
        for file_path in user_files[user_id].values():
            try:
                os.remove(file_path)
            except:
                pass
        del user_files[user_id]

async def main():
    """Главная функции"""
    # Создаем папку для временных файлов
    os.makedirs("temp_files", exist_ok=True)

    # Запускаем бота
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())