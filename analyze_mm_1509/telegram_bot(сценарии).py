import asyncio
import logging
from aiogram import Bot, Dispatcher, types
from aiogram.filters import Command
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
import os
from datetime import datetime
from docx import Document
from scenario_checker import check_meeting_scenario, safe_telegram_message

# Импортируем модуль для проверки критериев успеха
from success_criteria import check_success_criteria, find_meeting_file

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Конфигурация бота
BOT_TOKEN = "8079592721:AAGLaX7LwUPX0X5fr1SK-9IQnSvfP3Z96ws"
bot = Bot(token=BOT_TOKEN)
storage = MemoryStorage()
dp = Dispatcher(storage=storage)


# Состояния FSM
class UserStates(StatesGroup):
    waiting_for_scenario_check = State()
    waiting_for_success_criteria = State()  # Новое состояние для проверки критериев успеха
    asking_about_existing_file = State()  # Состояние для вопроса о ранее загруженной транскрибации


def escape_telegram_chars(text: str) -> str:
    """Экранирование специальных символов для Telegram"""
    escape_chars = ['_', '*', '[', ']', '(', ')', '~', '`', '>', '#', '+', '-', '=', '|', '{', '}', '.', '!']
    for char in escape_chars:
        text = text.replace(char, f'\\{char}')
    return text


def filter_facilitator_speech(text):
    """Фильтрует речь ведущего, определяя его по количеству предложений"""
    lines = text.split('\n')
    speaker_counts = {}

    # Подсчитываем количество предложений для каждого участника
    for line in lines:
        if ':' in line:
            speaker = line.split(':')[0].strip()
            if speaker and len(speaker) < 50:  # Исключаем слишком длинные имена
                speaker_counts[speaker] = speaker_counts.get(speaker, 0) + 1

    # Находим ведущего (того, кто говорит больше всего)
    if speaker_counts:
        facilitator = max(speaker_counts, key=speaker_counts.get)
        print(f"DEBUG: Определен ведущий: {facilitator} ({speaker_counts[facilitator]} предложений)")

        # Оставляем только речь ведущего
        filtered_lines = []
        for line in lines:
            if line.startswith(f"{facilitator}:"):
                filtered_lines.append(line)

        filtered_text = '\n'.join(filtered_lines)
        print(f"DEBUG: Оригинальный размер: {len(text)} символов")
        print(f"DEBUG: Размер после фильтрации: {len(filtered_text)} символов")
        return filtered_text

    return text


def read_docx_filtered(file_path):
    """Чтение и фильтрация текста из файла .docx или .txt"""
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.txt':
            # Читаем .txt файл
            with open(file_path, 'r', encoding='utf-8') as f:
                original_text = f.read()
        else:
            # Читаем .docx файл
            from docx import Document
            doc = Document(file_path)
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            original_text = '\n'.join(full_text)

        print(f"DEBUG: Оригинальный размер текста: {len(original_text)} символов")

        # Фильтруем речь ведущего
        filtered_text = filter_facilitator_speech(original_text)
        print(f"DEBUG: Размер после фильтрации: {len(filtered_text)} символов")

        return filtered_text
    except UnicodeDecodeError:
        # Пробуем другие кодировки для .txt
        try:
            with open(file_path, 'r', encoding='cp1251') as f:
                original_text = f.read()
            filtered_text = filter_facilitator_speech(original_text)
            return filtered_text
        except Exception as e:
            error_msg = f"Ошибка чтения файла: {str(e)}"
            print(escape_telegram_chars(error_msg))
            return ""
    except Exception as e:
        error_msg = f"Ошибка чтения файла: {str(e)}"
        print(escape_telegram_chars(error_msg))
        return ""


def save_filtered_text_to_file(filtered_text, filename_prefix="filtered_meeting"):
    """Сохраняет отфильтрованный текст в файл"""
    try:
        # Создаем папку для сохранения отфильтрованных текстов
        os.makedirs("filtered_texts", exist_ok=True)

        # Генерируем имя файла с временной меткой
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"filtered_texts/{filename_prefix}_{timestamp}.txt"

        # Сохраняем текст в файл
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(filtered_text)

        print(f"DEBUG: Отфильтрованный текст сохранен в файл: {filename}")
        return filename
    except Exception as e:
        print(f"Ошибка при сохранении отфильтрованного текста: {e}")
        return None


def save_result_to_docx(result_text, original_filename, result_type="scenario", user_id=None):
    """
    Сохраняет результат проверки в файл docx
    
    Args:
        result_text: текст результата проверки
        original_filename: имя исходного загруженного файла
        result_type: тип проверки ("scenario" или "competencies")
        user_id: ID пользователя для сохранения в его папку
    
    Returns:
        путь к сохраненному файлу или None при ошибке
    """
    try:
        # Пытаемся получить оригинальное имя файла из сохраненного файла
        display_filename = original_filename
        if user_id:
            user_folder = f"temp_files/{user_id}"
            original_filename_path = f"{user_folder}/original_filename.txt"
            if os.path.exists(original_filename_path):
                try:
                    with open(original_filename_path, 'r', encoding='utf-8') as f:
                        saved_filename = f.read().strip()
                        if saved_filename:
                            display_filename = saved_filename
                except Exception as e:
                    logger.warning(f"Не удалось прочитать оригинальное имя файла: {e}")
        
        # Извлекаем имя файла без расширения
        if display_filename:
            # Убираем расширение
            base_name = os.path.splitext(os.path.basename(display_filename))[0]
        else:
            base_name = "файл"
        
        # Формируем имя файла результата
        if result_type == "scenario":
            result_filename = f"проверка соотв. сценарию_{base_name}.docx"
        elif result_type == "competencies":
            result_filename = f"проверка компетенций модератора_{base_name}.docx"
        else:
            result_filename = f"результат проверки_{base_name}.docx"
        
        # Определяем папку для сохранения
        if user_id:
            save_folder = f"temp_files/{user_id}"
            os.makedirs(save_folder, exist_ok=True)
            file_path = os.path.join(save_folder, result_filename)
        else:
            save_folder = "temp_files"
            os.makedirs(save_folder, exist_ok=True)
            file_path = os.path.join(save_folder, result_filename)
        
        # Создаем документ Word
        doc = Document()
        
        # Добавляем заголовок
        if result_type == "scenario":
            doc.add_heading('Результат проверки соответствия сценарию', level=1)
        elif result_type == "competencies":
            doc.add_heading('Результат проверки компетенций модератора ММ', level=1)
        else:
            doc.add_heading('Результат проверки', level=1)
        
        # Добавляем информацию об исходном файле (используем оригинальное имя)
        if display_filename:
            doc.add_paragraph(f'Исходный файл: {os.path.basename(display_filename)}')
            doc.add_paragraph(f'Дата проверки: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph('')  # Пустая строка
        
        # Обрабатываем текст результата построчно для лучшей обработки markdown
        lines = result_text.split('\n')
        current_paragraph = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Пустая строка - завершаем текущий параграф
                if current_paragraph:
                    para_text = ' '.join(current_paragraph)
                    # Убираем markdown форматирование
                    para_text = para_text.replace('**', '').replace('*', '').replace('`', '').replace('_', '')
                    if para_text:
                        doc.add_paragraph(para_text)
                    current_paragraph = []
            elif line.startswith('#'):
                # Заголовок markdown
                if current_paragraph:
                    # Завершаем предыдущий параграф
                    para_text = ' '.join(current_paragraph)
                    para_text = para_text.replace('**', '').replace('*', '').replace('`', '').replace('_', '')
                    if para_text:
                        doc.add_paragraph(para_text)
                    current_paragraph = []
                
                # Обрабатываем заголовок
                level = len(line) - len(line.lstrip('#'))
                heading_text = line.lstrip('#').strip()
                # Убираем markdown форматирование из заголовка
                heading_text = heading_text.replace('**', '').replace('*', '').replace('`', '').replace('_', '')
                if heading_text:
                    doc.add_heading(heading_text, level=min(level, 3))
            else:
                # Обычная строка - добавляем к текущему параграфу
                current_paragraph.append(line)
        
        # Добавляем последний параграф, если он есть
        if current_paragraph:
            para_text = ' '.join(current_paragraph)
            para_text = para_text.replace('**', '').replace('*', '').replace('`', '').replace('_', '')
            if para_text:
                doc.add_paragraph(para_text)
        
        # Сохраняем документ
        doc.save(file_path)
        
        logger.info(f"✅ Результат проверки сохранен в файл: {file_path}")
        return file_path
        
    except Exception as e:
        logger.error(f"❌ Ошибка при сохранении результата в docx: {e}")
        return None


@dp.message(Command("start"))
async def cmd_start(message: types.Message):
    """Обработчик команды /start"""
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Проверить ММ на соответствие сценарию", callback_data="check_scenario")],
        [InlineKeyboardButton(text="✅ Проверить ММ на соответствие компетенциям модератора ММ",
                              callback_data="check_success_criteria")],
        [InlineKeyboardButton(text="ℹ️ Помощь", callback_data="help")]
    ])

    await message.answer(
        "👋 Добро пожаловать в бот анализа встреч!\n\n"
        "🤖 Я помогу вам проверить встречи на соответствие сценарию и компетенциям модератора.\n\n"
        "Выберите действие:",
        reply_markup=keyboard
    )


@dp.callback_query(lambda c: c.data == "check_success_criteria")
async def check_success_criteria_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик кнопки 'Проверить компетенции модератора ММ'"""
    try:
        user_id = callback.from_user.id
        
        # Проверяем, есть ли ранее загруженная транскрибация
        try:
            existing_file = find_meeting_file(user_id)
        except Exception as e:
            logger.error(f"Ошибка при поиске файла: {e}")
            existing_file = None
        
        if existing_file and os.path.exists(existing_file):
            # Если есть ранее загруженный файл - спрашиваем пользователя с указанием имени файла
            await state.set_state(UserStates.asking_about_existing_file)
            
            # Пытаемся получить оригинальное имя файла
            user_folder = f"temp_files/{user_id}"
            original_filename_path = f"{user_folder}/original_filename.txt"
            display_filename = os.path.basename(existing_file)
            
            if os.path.exists(original_filename_path):
                try:
                    with open(original_filename_path, 'r', encoding='utf-8') as f:
                        saved_filename = f.read().strip()
                        if saved_filename:
                            display_filename = saved_filename
                except Exception as e:
                    logger.warning(f"Не удалось прочитать оригинальное имя файла: {e}")
            
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="✅ Да, проверить по загруженной ранее", callback_data="use_existing_file")],
                [InlineKeyboardButton(text="📄 Нет, загрузить новую транскрибацию", callback_data="upload_new_file")],
                [InlineKeyboardButton(text="❌ Отменить", callback_data="cancel_operation")]
            ])
            
            # Экранируем специальные символы Markdown в имени файла
            safe_filename = escape_telegram_chars(display_filename)
            
            await callback.message.answer(
                f"🏆 **Проверка компетенций модератора ММ**\n\n"
                f"📋 Проверить ранее загруженную транскрибацию \"{safe_filename}\"?",
                reply_markup=keyboard,
                parse_mode="Markdown"
            )
        else:
            # Если нет ранее загруженного файла - сразу просим загрузить новый
            await state.set_state(UserStates.waiting_for_success_criteria)
            
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="❌ Отменить", callback_data="cancel_operation")]
            ])
            
            await callback.message.answer(
                "🏆 **Проверка компетенций модератора ММ**\n\n"
                "📄 Отправьте файл с текстом встречи (.docx или .txt) для проверки соответствия компетенциям модератора ММ.\n\n"
                "Для отмены нажмите кнопку ниже или отправьте /cancel",
                reply_markup=keyboard,
                parse_mode="Markdown"
            )
        
        await callback.answer()
        
    except Exception as e:
        logger.error(f"Ошибка в check_success_criteria_handler: {e}", exc_info=True)
        try:
            await callback.message.answer(
                "❌ Произошла ошибка при обработке запроса. Попробуйте еще раз или обратитесь к администратору."
            )
            await callback.answer()
        except Exception:
            pass


@dp.callback_query(lambda c: c.data == "use_existing_file")
async def use_existing_file_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик кнопки 'Да, проверить по загруженной ранее'"""
    user_id = callback.from_user.id
    
    # Ищем ранее загруженный файл
    file_path = find_meeting_file(user_id)
    
    if not file_path:
        await callback.message.answer(
            "❌ Ранее загруженная транскрибация не найдена. Пожалуйста, загрузите файл."
        )
        await state.set_state(UserStates.waiting_for_success_criteria)
        await callback.answer()
        return
    
    await callback.message.answer(
        f"✅ Использую ранее загруженную транскрибацию: {os.path.basename(file_path)}\n\n"
        "🔄 Анализирую текст встречи..."
    )
    
    # Показываем индикатор обработки
    processing_msg = await callback.message.answer("⏳ Обрабатываю...")
    
    try:
        # Вызываем функцию проверки критериев успеха
        success_result = await check_success_criteria(file_path=file_path, user_id=user_id)
        
        # Сохраняем результат проверки в файл docx
        original_filename = os.path.basename(file_path)
        result_file_path = save_result_to_docx(
            result_text=success_result,
            original_filename=original_filename,
            result_type="competencies",
            user_id=user_id
        )
        if result_file_path:
            logger.info(f"📄 Результат проверки компетенций сохранен: {result_file_path}")
        
        # Удаляем индикатор обработки
        await bot.delete_message(chat_id=callback.message.chat.id, message_id=processing_msg.message_id)
        
        # Безопасная отправка результата пользователю
        message_parts = safe_telegram_message(success_result)
        
        for i, part in enumerate(message_parts):
            if i == 0:
                # Первую часть отправляем с заголовком
                await callback.message.answer(
                    f"🏆 **Результат проверки компетенций модератора ММ:**\n\n{part}",
                    parse_mode="Markdown"
                )
            else:
                # Остальные части просто как текст
                await callback.message.answer(part)
        
        # Отправляем кнопки меню
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_to_main"),
             InlineKeyboardButton(text="🔄 Проверить другой файл", callback_data="check_success_criteria")]
        ])
        
        await callback.message.answer(
            "✅ Проверка завершена!\n\n"
            "Вы можете:\n"
            "• Вернуться в главное меню\n"
            "• Проверить другой файл",
            reply_markup=keyboard
        )
        
        # Сбрасываем состояние
        await state.clear()
        
    except Exception as e:
        error_msg = f"❌ Ошибка при проверке компетенций: {str(e)}"
        logger.error(error_msg)
        await callback.message.answer(error_msg)
        await state.clear()
    
    await callback.answer()


@dp.callback_query(lambda c: c.data == "upload_new_file")
async def upload_new_file_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик кнопки 'Нет, загрузить новую транскрибацию'"""
    await state.set_state(UserStates.waiting_for_success_criteria)
    
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="❌ Отменить", callback_data="cancel_operation")]
    ])
    
    await callback.message.answer(
        "📄 **Загрузка новой транскрибации**\n\n"
        "Отправьте файл с текстом встречи (.docx или .txt) для проверки соответствия компетенциям модератора ММ.\n\n"
        "Для отмены нажмите кнопку ниже или отправьте /cancel",
        reply_markup=keyboard,
        parse_mode="Markdown"
    )
    await callback.answer()


@dp.message(UserStates.waiting_for_success_criteria)
async def handle_success_criteria_file(message: types.Message, state: FSMContext):
    """Обработчик файла для проверки критериев успеха"""
    user_id = message.from_user.id

    if not message.document:
        await message.answer("❌ Пожалуйста, отправьте файл с текстом встречи.")
        return

    file_name = message.document.file_name
    if not file_name.lower().endswith(('.docx', '.txt')):
        await message.answer("❌ Пожалуйста, отправьте файл в формате .docx или .txt")
        return

    await message.answer("📥 Загружаю файл для проверки компетенций модератора ММ...")

    try:
        # Скачиваем файл
        file_info = await bot.get_file(message.document.file_id)
        downloaded_file = await bot.download_file(file_info.file_path)

        # Создаем папку для пользователя
        user_folder = f"temp_files/{user_id}"
        os.makedirs(user_folder, exist_ok=True)

        # Сохраняем файл с фиксированным именем для последующего использования
        # Определяем расширение файла
        file_ext = os.path.splitext(file_name)[1].lower()
        if file_ext == '.txt':
            file_path = f"{user_folder}/trans.txt"
        else:
            file_path = f"{user_folder}/trans.docx"
        
        # Сохраняем оригинальное имя файла для отображения пользователю
        original_filename_path = f"{user_folder}/original_filename.txt"
        with open(original_filename_path, 'w', encoding='utf-8') as f:
            f.write(file_name)
        
        # Если файл уже существует, заменяем его
        with open(file_path, 'wb') as f:
            f.write(downloaded_file.read())

        await message.answer("✅ Файл загружен! Проверяю встречу на соответствие компетенциям модератора ММ...")

        # Показываем индикатор обработки
        processing_msg = await message.answer("🔄 Анализирую текст встречи...")

        # Вызываем функцию проверки критериев успеха
        success_result = await check_success_criteria(file_path)

        # Сохраняем результат проверки в файл docx
        result_file_path = save_result_to_docx(
            result_text=success_result,
            original_filename=file_name,
            result_type="competencies",
            user_id=user_id
        )
        if result_file_path:
            logger.info(f"📄 Результат проверки компетенций сохранен: {result_file_path}")

        # Удаляем индикатор обработки
        await bot.delete_message(chat_id=message.chat.id, message_id=processing_msg.message_id)

        # Безопасная отправка результата пользователю
        message_parts = safe_telegram_message(success_result)

        for i, part in enumerate(message_parts):
            if i == 0:
                # Первую часть отправляем с заголовком
                await message.answer(f"🏆 **Результат проверки компетенций модератора ММ:**\n\n{part}",
                                     parse_mode="Markdown")
            else:
                # Остальные части просто как текст
                await message.answer(part)

        # Отправляем кнопку меню
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_to_main"),
             InlineKeyboardButton(text="🔄 Проверить другой файл", callback_data="check_success_criteria")]
        ])

        await message.answer(
            "✅ Проверка завершена!\n\n"
            "Вы можете:\n"
            "• Вернуться в главное меню\n"
            "• Проверить другой файл на компетенции модератора",
            reply_markup=keyboard
        )
        
        # Сбрасываем состояние
        await state.clear()

    except Exception as e:
        error_msg = f"❌ Ошибка при проверке компетенций модератора: {str(e)}"
        logging.error(f"Ошибка в handle_success_criteria_file: {e}")

        # Отправляем сообщение об ошибке
        await message.answer(error_msg)

        # Предлагаем вернуться в меню
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_to_main")]
        ])
        await message.answer("Произошла ошибка. Вернитесь в главное меню:", reply_markup=keyboard)

        await state.clear()


@dp.callback_query(lambda c: c.data == "check_scenario")
async def check_scenario_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик кнопки 'Проверить ММ на соответствие сценарию' — сразу переход на универсальный сценарий"""
    await state.update_data(scenario_type="scenario_universal")
    await state.set_state(UserStates.waiting_for_scenario_check)

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="❌ Отменить", callback_data="cancel_operation")]
    ])

    await callback.message.answer(
        "📄 Для проверки на соответствие *универсальному сценарию* отправьте файл с текстом встречи (.docx или .txt)\n\n"
        "Для отмены нажмите кнопку ниже или отправьте /cancel",
        reply_markup=keyboard,
        parse_mode="Markdown"
    )
    await callback.answer()


@dp.callback_query(lambda c: c.data.startswith("scenario_"))
async def scenario_type_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик выбора типа ММ для проверки сценария"""
    scenario_type = callback.data

    # Сохраняем выбранный тип сценария в состоянии
    await state.update_data(scenario_type=scenario_type)
    await state.set_state(UserStates.waiting_for_scenario_check)

    # Определяем название типа для отображения
    scenario_names = {
        "scenario_online": "ММ онлайн торговля",
        "scenario_first_meetings": "ММ первые встречи",
        "scenario_first_month": "ММ первый месяц",
        "scenario_my_meetings": "ММ мои встречи",
        "scenario_universal": "ММ универсальный сценарий"
    }

    scenario_name = scenario_names.get(scenario_type, "выбранный тип ММ")

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="❌ Отменить", callback_data="cancel_operation")]
    ])

    await callback.message.answer(
        f"📄 Для проверки *{scenario_name}* на соответствие сценарию отправьте файл с текстом встречи (.docx или .txt)\n\n"
        "Для отмена нажмите кнопку ниже или отправьте /cancel",
        reply_markup=keyboard,
        parse_mode="Markdown"
    )
    await callback.answer()


@dp.message(UserStates.waiting_for_scenario_check)
async def handle_scenario_check_file(message: types.Message, state: FSMContext):
    """Обработчик файла для проверки сценария"""
    user_id = message.from_user.id

    if not message.document:
        await message.answer("❌ Пожалуйста, отправьте файл с текстом встречи.")
        return

    file_name = message.document.file_name
    if not file_name.lower().endswith(('.docx', '.txt')):
        await message.answer("❌ Пожалуйста, отправьте файл в формате .docx или .txt")
        return

    # Получаем тип сценария из состояния (по умолчанию — универсальный)
    user_data = await state.get_data()
    scenario_type = user_data.get('scenario_type', 'scenario_universal')

    scenario_names = {
        "scenario_online": "ММ онлайн торговля",
        "scenario_first_meetings": "ММ первые встречи",
        "scenario_first_month": "ММ первый месяц",
        "scenario_my_meetings": "ММ мои встречи",
        "scenario_universal": "ММ универсальный сценарий"
    }

    scenario_name = scenario_names.get(scenario_type, "ММ")

    await message.answer(f"📥 Загружаю файл для проверки {scenario_name}...")

    try:
        # Скачиваем файл
        file_info = await bot.get_file(message.document.file_id)
        downloaded_file = await bot.download_file(file_info.file_path)

        # Создаем папку для пользователя
        user_folder = f"temp_files/{user_id}"
        os.makedirs(user_folder, exist_ok=True)

        # Сохраняем файл с фиксированным именем для последующего использования при проверке компетенций
        # Определяем расширение файла
        file_ext = os.path.splitext(file_name)[1].lower()
        if file_ext == '.txt':
            file_path = f"{user_folder}/trans.txt"
        else:
            file_path = f"{user_folder}/trans.docx"
        
        # Сохраняем оригинальное имя файла для отображения пользователю
        original_filename_path = f"{user_folder}/original_filename.txt"
        with open(original_filename_path, 'w', encoding='utf-8') as f:
            f.write(file_name)
        
        # Сохраняем файл (заменяем, если уже существует)
        with open(file_path, 'wb') as f:
            f.write(downloaded_file.read())

        # Читаем и фильтруем текст
        print("\n" + "=" * 80)
        print("DEBUG: ОБРАБОТКА ТЕКСТА ВСТРЕЧИ")
        print("=" * 80)

        # Читаем оригинальный текст (поддерживаем .docx и .txt)
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    original_text = f.read()
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='cp1251') as f:
                    original_text = f.read()
        else:
            from docx import Document
            doc = Document(file_path)
            original_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])

        print(f"ОРИГИНАЛЬНЫЙ ТЕКСТ ({len(original_text)} символов):")
        print("-" * 40)
        print(original_text[:1000] + "..." if len(original_text) > 1000 else original_text)
        print("-" * 40)

        # Фильтруем текст (оставляем только речь ведущего)
        filtered_text = read_docx_filtered(file_path)

        print(f"ОТФИЛЬТРОВАННЫЙ ТЕКСТ ({len(filtered_text)} символов):")
        print("-" * 40)
        print(filtered_text[:1000] + "..." if len(filtered_text) > 1000 else filtered_text)
        print("-" * 40)
        print(f"Сокращение текста: {len(original_text)} -> {len(filtered_text)} символов "
              f"({(1 - len(filtered_text) / len(original_text)) * 100:.1f}% сокращено)")

        # СОХРАНЯЕМ ОТФИЛЬТРОВАННЫЙ ТЕКСТ В ФАЙЛ
        saved_file_path = save_filtered_text_to_file(filtered_text, f"scenario_{scenario_type}")
        if saved_file_path:
            print(f"DEBUG: Весь отфильтрованный текст сохранен в: {saved_file_path}")

        print("=" * 80 + "\n")

        await message.answer(f"✅ Файл загружен! Проверяю {scenario_name} на соответствие сценарию...")

        # Показываем индикатор обработки
        processing_msg = await message.answer("🔄 Анализирую текст встречи...")

        # СОЗДАЕМ ВРЕМЕННЫЙ ФАЙЛ С ОТФИЛЬТРОВАННЫМ ТЕКСТОМ ДЛЯ ОТПРАВКИ В GIGACHAT
        filtered_file_path = f"{user_folder}/filtered_for_gigachat.docx"

        # Создаем новый DOCX файл с отфильтрованным текстом
        from docx import Document as DocxDocument
        filtered_doc = DocxDocument()
        for line in filtered_text.split('\n'):
            if line.strip():
                filtered_doc.add_paragraph(line)
        filtered_doc.save(filtered_file_path)

        # Вызываем функцию проверки сценария с ОТФИЛЬТРОВАННЫМ файлом
        scenario_result = await check_meeting_scenario(filtered_file_path, scenario_type)

        # Сохраняем результат проверки в файл docx
        result_file_path = save_result_to_docx(
            result_text=scenario_result,
            original_filename=file_name,
            result_type="scenario",
            user_id=user_id
        )
        if result_file_path:
            logger.info(f"📄 Результат проверки сценария сохранен: {result_file_path}")

        # Удаляем индикатор обработки
        await bot.delete_message(chat_id=message.chat.id, message_id=processing_msg.message_id)

        # Безопасная отправка результата пользователю
        message_parts = safe_telegram_message(scenario_result)

        for i, part in enumerate(message_parts):
            if i == 0:
                # Первую часть отправляем с заголовком
                await message.answer(part, parse_mode="Markdown")
            else:
                # Остальные части просто как текст
                await message.answer(part)

        # Отправляем кнопку меню
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_to_main"),
             InlineKeyboardButton(text="🔄 Проверить другой сценарий", callback_data="check_scenario")]
        ])

        await message.answer(
            "✅ Проверка завершена!\n\n"
            "Вы можете:\n"
            "• Вернуться в главное меню\n"
            "• Проверить другой сценарий",
            reply_markup=keyboard
        )

        # Удаляем только временный отфильтрованный файл
        # Оригинальный файл (file_path) НЕ удаляем - он нужен для проверки компетенций
        try:
            os.remove(filtered_file_path)
        except Exception as e:
            logging.warning(f"Не удалось удалить временный файл: {e}")

        await state.clear()

    except Exception as e:
        error_msg = f"❌ Ошибка при проверке сценария: {str(e)}"
        logging.error(f"Ошибка в handle_scenario_check_file: {e}")

        # Отправляем сообщение об ошибке
        await message.answer(error_msg)

        # Предлагаем вернуться в меню
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🏠 Главное меню", callback_data="back_to_main")]
        ])
        await message.answer("Произошла ошибка. Вернитесь в главное меню:", reply_markup=keyboard)

        await state.clear()


@dp.callback_query(lambda c: c.data == "help")
async def help_handler(callback: types.CallbackQuery):
    """Обработчик кнопки 'Помощь'"""
    help_text = """
ℹ️ **ПОМОЩЬ ПО ИСПОЛЬЗОВАНИЮ БОТА**

✅ **Проверить ММ на соответствие сценарию:**
• Выберите тип ММ: онлайн торговля, первые встречи, первый месяц, мои встречи или универсальный сценарий
• Загрузите файл с текстом встречи (trans.docx)
• Получите анализ соответствия проведенной встречи запланированному сценарию

🏆 **Проверить на соответствие компетенциям модератора ММ:**
• Загрузите файл с текстом встречи (trans.docx)
• Получите проверку соответствия копметенциям модератора ММ

📁 **Форматы файлов:**
• trans.docx - текст встречи с временными метками

⏱️ **Время обработки:**
• Проверка сценария: 10-20 секунд
• Проверка компетенций: 10-20 секунд

🔧 **При возникновении проблем:**
• Проверьте формат файлов
• Убедитесь в наличии интернета
• Попробуйте загрузить файлы заново
"""

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🔙 Назад", callback_data="back_to_main")]
    ])

    await callback.message.edit_text(help_text, reply_markup=keyboard, parse_mode="Markdown")
    await callback.answer()


@dp.callback_query(lambda c: c.data == "back_to_main")
async def back_to_main_handler(callback: types.CallbackQuery):
    """Обработчик кнопки 'Назад'"""
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Проверить ММ на соответствие сценарию", callback_data="check_scenario")],
        [InlineKeyboardButton(text="✅ Проверить на соответствие компетенциям модератора ММ",
                              callback_data="check_success_criteria")],
        [InlineKeyboardButton(text="ℹ️ Помощь", callback_data="help")]
    ])

    await callback.message.edit_text(
        "👋 Добро пожаловать в бот анализа встреч!\n\n"
        "🤖 Я помогу вам проверить встречи на соответствие сценарию и компетенциям модератора.\n\n"
        "Выберите действие:",
        reply_markup=keyboard
    )
    await callback.answer()


@dp.callback_query(lambda c: c.data == "cancel_operation")
async def cancel_operation_handler(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик отмены операции с очисткой временных файлов"""
    user_id = callback.from_user.id
    user_folder = f"temp_files/{user_id}"

    # Очищаем временные файлы
    try:
        if os.path.exists(user_folder):
            import shutil
            shutil.rmtree(user_folder)
    except Exception as e:
        logging.warning(f"Не удалось очистить временные файлы при отмене: {e}")

    await state.clear()

    # Возвращаем в главное меню
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Проверить ММ на соответствие сценарию", callback_data="check_scenario")],
        [InlineKeyboardButton(text="✅ Проверить на соответствие компетенциям модератора ММ",
                              callback_data="check_success_criteria")],
        [InlineKeyboardButton(text="ℹ️ Помощь", callback_data="help")]
    ])

    await callback.message.edit_text(
        "❌ Операция отменена. Выберите новое действие:",
        reply_markup=keyboard
    )
    await callback.answer()


async def main():
    """Главная функция"""
    # Создаем папки для временных файлов
    os.makedirs("temp_files", exist_ok=True)
    os.makedirs("filtered_texts", exist_ok=True)

    # Запускаем бота
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())